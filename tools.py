import json
import csv
import io
from datetime import datetime
from pathlib import Path
from langchain_core.tools import tool
from pydantic import BaseModel
from schemas import PRDDocument


# ── PDF helper ────────────────────────────────────────────────────────────────

def _md_to_pdf(md_content: str, pdf_path: Path) -> tuple[bool, str]:
    """
    Convert a Markdown string to a styled PDF.
    Returns (success, error_message).
    Falls back gracefully — JSON/MD saves are never blocked by PDF failure.
    Requires: pip install weasyprint markdown
    """
    try:
        import markdown as _md
    except ImportError as e:
        return False, f"markdown not installed: {e}"

    try:
        from weasyprint import HTML as _HTML
    except Exception as e:
        return False, f"weasyprint not available: {e}"

    try:
        html_body = _md.markdown(
            md_content,
            extensions=["tables", "fenced_code", "nl2br"],
        )
        html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
  @page {{ margin: 30mm 20mm; }}
  body  {{ font-family: Arial, sans-serif; font-size: 10.5pt;
           line-height: 1.55; color: #1a1a1a; }}
  h1   {{ font-size: 18pt; border-bottom: 3px solid #F5E003;
          padding-bottom: 8px; margin-top: 0; }}
  h2   {{ font-size: 13pt; color: #222; border-bottom: 1px solid #ddd;
          padding-bottom: 4px; margin-top: 22px; }}
  h3   {{ font-size: 11pt; color: #444; margin-top: 16px; }}
  table {{ border-collapse: collapse; width: 100%;
           margin: 10px 0; font-size: 9.5pt; page-break-inside: auto; }}
  th   {{ background: #F5E003; color: #1a1a1a; padding: 5px 9px;
          text-align: left; font-weight: 600; }}
  td   {{ border: 1px solid #ddd; padding: 5px 9px; }}
  tr   {{ page-break-inside: avoid; }}
  tr:nth-child(even) {{ background: #f8f8f8; }}
  code {{ background: #f3f3f3; padding: 1px 4px; border-radius: 3px;
          font-family: "Courier New", monospace; font-size: 9pt; }}
  pre  {{ background: #f3f3f3; padding: 10px; border-radius: 4px;
          font-size: 8.5pt; white-space: pre-wrap; }}
  hr   {{ border: none; border-top: 1px solid #e0e0e0; margin: 16px 0; }}
  ul, ol {{ margin: 6px 0; padding-left: 22px; }}
  li   {{ margin-bottom: 3px; }}
</style>
</head><body>{html_body}</body></html>"""
        _HTML(string=html).write_pdf(str(pdf_path))
        return True, ""
    except Exception as e:
        return False, str(e)


class _NoInput(BaseModel):
    @classmethod
    def model_json_schema(cls, **kwargs) -> dict:
        return {"type": "object", "properties": {}}
    @classmethod
    def schema(cls, **kwargs) -> dict:
        return {"type": "object", "properties": {}}

SDLC_DIR           = Path(__file__).parent.parent
PRODUCT_CTX_FILE   = Path(__file__).parent / "product_context.md"
OUTPUT_DIR         = Path(__file__).parent / "output"
try:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
except OSError as e:
    raise RuntimeError(f"Could not create output directory {OUTPUT_DIR}: {e}")


# ── File readers ──────────────────────────────────────────────────────────────

def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages  = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(p.strip() for p in pages if p.strip())
    except Exception as e:
        return f"[PDF read error: {e}]"


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
        doc  = Document(str(path))
        text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        # also pull text from tables
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + " | ".join(cell.text.strip() for cell in row.cells)
        return text
    except Exception as e:
        return f"[DOCX read error: {e}]"


def _read_csv(path: Path) -> str:
    try:
        import pandas as pd
        df = pd.read_csv(str(path))
        return f"CSV with {len(df)} rows and {len(df.columns)} columns.\n\nColumns: {list(df.columns)}\n\nFirst 20 rows:\n{df.head(20).to_string(index=False)}"
    except Exception as e:
        return f"[CSV read error: {e}]"


def _read_txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        return f"[TXT read error: {e}]"


# ── PRD markdown renderer ─────────────────────────────────────────────────────

def _prd_to_markdown(prd: PRDDocument) -> str:
    def checklist(items): return "\n".join(f"- [ ] {i}" for i in items) or "- None"
    def bullets(items):   return "\n".join(f"- {i}" for i in items) or "- None"

    stories_md = ""
    for s in prd.user_stories:
        ac     = "\n".join(f"  - [ ] {c.criterion}" for c in s.acceptance_criteria)
        labels = ", ".join(f"`{l}`" for l in s.labels) if s.labels else "—"
        stories_md += f"""
### {s.id}: {s.title}
**As a** {s.as_a},
**I want to** {s.i_want_to},
**So that** {s.so_that}.

| Priority | Story Points | Labels |
|----------|-------------|--------|
| {s.priority.value} | {s.story_points} | {labels} |

**Acceptance Criteria:**
{ac}

---"""

    goals_rows        = "\n".join(f"| {g.goal} | {g.metric} | {g.baseline} | {g.target} |" for g in prd.goals)
    personas_rows     = "\n".join(f"| {p.name} | {p.description} | {p.key_need} |" for p in prd.personas)
    fr_rows           = "\n".join(f"| {r.id} | {r.requirement} | {r.priority.value} |" for r in prd.functional_requirements)
    nfr_rows          = "\n".join(f"| {n.id} | {n.category.value} | {n.requirement} | {n.target} |" for n in prd.non_functional_requirements)
    edge_rows         = "\n".join(f"| {e.scenario} | {e.expected_behaviour} |" for e in prd.edge_cases)
    risk_rows         = "\n".join(f"| {r.risk} | {r.impact} | {r.mitigation} |" for r in prd.risks)
    stakeholder_rows  = "\n".join(f"| {s.name} | {s.role} | {s.email or '—'} |" for s in prd.stakeholders)

    return f"""# PRD: {prd.title}

| Field | Value |
|-------|-------|
| **Version** | {prd.version} |
| **Date** | {prd.date} |
| **Status** | {prd.status.value} |
| **Type** | {prd.prd_type.value} |
| **Author** | {prd.author} |

---

## 1. Problem Statement
{prd.problem_statement}

## 2. Business Value
{prd.business_value}

## 3. Stakeholders
| Name | Role | Email |
|------|------|-------|
{stakeholder_rows}

## 4. Goals & Success Metrics
| Goal | Metric | Baseline | Target |
|------|--------|----------|--------|
{goals_rows}

## 5. Scope
### In Scope (v1)
{bullets(prd.in_scope)}

### Out of Scope
{bullets(prd.out_of_scope)}

## 6. Users & Personas
| Persona | Description | Key Need |
|---------|-------------|----------|
{personas_rows}

## 7. Functional Requirements
| ID | Requirement | Priority |
|----|-------------|----------|
{fr_rows}

## 8. Non-Functional Requirements
| ID | Category | Requirement | Target |
|----|----------|-------------|--------|
{nfr_rows if nfr_rows else "| — | — | No NFRs specified | — |"}

## 9. User Stories
{stories_md}

## 10. Edge Cases & Error Scenarios
| Scenario | Expected Behaviour |
|----------|--------------------|
{edge_rows}

## 11. Risks
| Risk | Impact | Mitigation |
|------|--------|------------|
{risk_rows if risk_rows else "| — | — | No risks identified |"}

## 12. Technical Constraints
{bullets(prd.technical_constraints)}

## 13. Dependencies
{bullets(prd.dependencies)}

## 14. Assumptions
{bullets(prd.assumptions)}

## 15. Open Questions
{checklist(prd.open_questions)}
"""


# ── Tools ─────────────────────────────────────────────────────────────────────

@tool
def get_current_date(query: str) -> str:
    """Get today's date in YYYY-MM-DD format. Use this to populate PRD metadata.
    Pass any short string (like 'today') for the query parameter — it is ignored."""
    return datetime.now().strftime("%Y-%m-%d")


@tool
def read_input_file(filepath: str) -> str:
    """
    Read and extract text from an input file provided by the user.
    Supports: PDF (.pdf), Word document (.docx), CSV (.csv),
    plain text (.txt), and Teams conversation export (.txt).

    After reading, extract as much PRD-relevant information as possible
    from the content: problem, personas, requirements, constraints, etc.
    Then only ask the user for what is genuinely missing.

    Args:
        filepath: Absolute or relative path to the file.
    """
    path = Path(filepath.strip().strip('"').strip("'"))

    if not path.exists():
        return json.dumps({
            "error": f"File not found: {filepath}",
            "hint": "Make sure the path is correct. On Windows use forward slashes or double backslashes."
        })

    suffix = path.suffix.lower()
    size_kb = round(path.stat().st_size / 1024, 1)

    if suffix == ".pdf":
        content = _read_pdf(path)
        file_type = "PDF document"
    elif suffix == ".docx":
        content = _read_docx(path)
        file_type = "Word document"
    elif suffix == ".csv":
        content = _read_csv(path)
        file_type = "CSV file"
    elif suffix in (".txt", ".md", ".text"):
        content = _read_txt(path)
        file_type = "Text file (may contain Teams conversation or plain description)"
    else:
        return json.dumps({
            "error": f"Unsupported file type: {suffix}",
            "supported": [".pdf", ".docx", ".csv", ".txt", ".md"]
        })

    if not content.strip():
        return json.dumps({"error": "File appears to be empty or could not be parsed."})

    # Truncate very large files to avoid context overflow (keep first ~32000 chars)
    truncated = False
    if len(content) > 32000:
        content   = content[:32000]
        truncated = True

    instruction = (
        "Extract every PRD-relevant detail from this content: problem statement, "
        "business context, user types, requirements, constraints, stakeholders, goals, "
        "timelines, and any open questions. Then identify ONLY what is missing or unclear "
        "and ask the user for just those gaps — one question at a time."
    )
    if truncated:
        instruction += (
            " NOTE: This file was truncated to 32KB — the original was larger. "
            "Some content near the end of the file may be missing. "
            "If key sections seem incomplete, ask the user to provide the missing details."
        )

    return json.dumps({
        "file_type":   file_type,
        "filename":    path.name,
        "size_kb":     size_kb,
        "truncated":   truncated,
        "content":     content,
        "instruction": instruction,
    })


@tool
def list_existing_prds(query: str) -> str:
    """
    List all previously saved PRD files so the user can load and revise one.
    Returns file names, versions, and last modified timestamps.
    Pass any short string (like 'list') for the query parameter — it is ignored.
    """
    files = sorted(OUTPUT_DIR.glob("*.json"), key=lambda f: f.stat().st_mtime, reverse=True)
    if not files:
        return "No existing PRDs found. Starting fresh."
    result = []
    for f in files:
        try:
            data = json.loads(f.read_text(encoding="utf-8"))
            result.append({
                "filename": f.name,
                "title":    data.get("title", "Unknown"),
                "version":  data.get("version", "1.0"),
                "status":   data.get("status", "Unknown"),
                "type":     data.get("prd_type", "Unknown"),
                "modified": datetime.fromtimestamp(f.stat().st_mtime).strftime("%Y-%m-%d %H:%M"),
            })
        except Exception:
            result.append({"filename": f.name, "error": "Could not parse"})
    return json.dumps(result, indent=2)


@tool
def load_prd(filename: str) -> str:
    """
    Load an existing PRD by filename for revision.
    Args:
        filename: The .json filename from list_existing_prds output.
    """
    path = OUTPUT_DIR / filename
    if not path.exists():
        return json.dumps({"error": f"File not found: {filename}"})
    return path.read_text(encoding="utf-8")


@tool
def save_prd(prd_json: str) -> str:
    """
    Validate and save the completed PRD as both .json (machine-readable)
    and .md (human-readable). Call ONLY after the user confirms.
    Args:
        prd_json: A JSON string representing the full PRDDocument schema.
    """
    try:
        data = json.loads(prd_json)
    except json.JSONDecodeError as e:
        return json.dumps({
            "error": f"PRD JSON is not valid JSON: {e}",
            "hint":  "Check for unescaped quotes or trailing commas in the string you passed."
        })
    try:
        prd  = PRDDocument(**data)
    except Exception as e:
        return json.dumps({
            "error": f"PRD validation failed: {e}",
            "hint":  (
                "Re-read the PRDDocument schema in the system prompt. "
                "Common causes: wrong field name ('problem' instead of 'problem_statement', "
                "'constraints' instead of 'technical_constraints', 'nfrs' instead of "
                "'non_functional_requirements'), stakeholders/goals/FRs/NFRs/risks/edge_cases "
                "as strings instead of structured objects, goals missing 'baseline' field, "
                "story_points not in Fibonacci set {1,2,3,5,8,13}, missing prd_type, "
                "invalid NFR category (must be Performance/Security/Scalability/Accessibility/Reliability), "
                "or fewer than 3 FRs / 3 stories / 2 edge cases. Fix the JSON and call save_prd again."
            )
        })

    timestamp  = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = prd.title.lower().replace(" ", "_")[:40]
    base       = f"{timestamp}_{safe_title}_v{prd.version.replace('.', '_')}"

    json_path  = OUTPUT_DIR / f"{base}.json"
    md_path    = OUTPUT_DIR / f"{base}.md"
    pdf_path   = OUTPUT_DIR / f"{base}.pdf"

    md_content = _prd_to_markdown(prd)
    json_path.write_text(prd.model_dump_json(indent=2), encoding="utf-8")
    md_path.write_text(md_content, encoding="utf-8")

    try:
        pdf_ok, pdf_err = _md_to_pdf(md_content, pdf_path)
    except Exception as e:
        pdf_ok, pdf_err = False, str(e)

    return json.dumps({
        "saved":        True,
        "title":        prd.title,
        "version":      prd.version,
        "json_path":    str(json_path),
        "md_path":      str(md_path),
        "pdf_path":     str(pdf_path) if pdf_ok else None,
        "pdf_error":    pdf_err if not pdf_ok else None,
        "stories":      len(prd.user_stories),
        "requirements": len(prd.functional_requirements),
    })


@tool
def load_product_context(query: str) -> str:
    """
    Load the TestOps product context document. Call this ONCE at the start
    of every conversation so you understand:
      - What TestOps is (A/B testing platform for AB InBev)
      - Tech stack (React + FastAPI + PostgreSQL + Azure DevOps)
      - Existing pages, components, domain terms (KPI, uplift, BEES, etc.)

    Pass any short string (like 'load') — it is ignored.
    """
    if not PRODUCT_CTX_FILE.exists():
        return json.dumps({
            "error": f"product_context.md not found at {PRODUCT_CTX_FILE}",
            "hint":  "Create this file with TestOps product knowledge."
        })
    try:
        content = PRODUCT_CTX_FILE.read_text(encoding="utf-8")
        return json.dumps({"loaded": True, "content": content})
    except Exception as e:
        return json.dumps({"error": f"Could not read product_context.md: {e}"})


# ═════════════════════════════════════════════════════════════════════════════
# SHAREPOINT TOOLS
# ═════════════════════════════════════════════════════════════════════════════

# Cache the last folder listing so read_sharepoint_file can auto-resolve
# numeric display indices (e.g. "3") to real Graph item IDs.
_sp_folder_cache: list[dict] = []


@tool
def list_sharepoint_folder(folder_path: str) -> str:
    """
    List files and sub-folders inside a SharePoint folder.

    Args:
        folder_path: Site-relative path, e.g. "Documents/Briefs".
                     Pass an empty string "" or "/" to list the root library.

    Returns JSON with each item's id, name, type (file/folder), extension,
    size_kb, and modified date.  Use the item `id` when calling
    read_sharepoint_file().

    Only available when SHAREPOINT_SITE_URL, SHAREPOINT_TENANT_ID,
    SHAREPOINT_CLIENT_ID, and SHAREPOINT_CLIENT_SECRET are set in the .env.
    """
    try:
        import sharepoint_client as sp
    except ImportError as e:
        return json.dumps({"error": f"sharepoint_client import failed: {e}"})

    if not sp.is_configured():
        return json.dumps({
            "error": "SharePoint not configured.",
            "hint":  (
                "Add SHAREPOINT_TENANT_ID, SHAREPOINT_CLIENT_ID, "
                "SHAREPOINT_CLIENT_SECRET, SHAREPOINT_SITE_URL to SDLC/.env"
            ),
        })

    try:
        items = sp.list_folder(folder_path)
    except Exception as e:
        return json.dumps({"error": str(e)})

    if not items:
        return json.dumps({
            "folder":  folder_path or "(root)",
            "items":   [],
            "message": "Folder is empty or does not exist.",
        })

    files   = [i for i in items if i["type"] == "file"]
    folders = [i for i in items if i["type"] == "folder"]

    supported_exts = {".pdf", ".docx", ".xlsx", ".csv", ".txt", ".md"}
    readable = [f for f in files if f["extension"] in supported_exts]

    # Cache readable files so read_sharepoint_file can resolve numeric picks
    global _sp_folder_cache
    _sp_folder_cache = readable

    return json.dumps({
        "folder":           folder_path or "(root)",
        "total_items":      len(items),
        "readable_files":   readable,
        "other_files":      [f for f in files if f["extension"] not in supported_exts],
        "sub_folders":      folders,
        "instruction":      (
            "Show the user the readable_files list (name + size + modified). "
            "Ask which file they want to use as input for the PRD. "
            "When they choose, call read_sharepoint_file() using the 'id' field "
            "from that file's entry in readable_files — this is a long alphanumeric "
            "string (e.g. '01ABC...'). NEVER pass a display index like '1' or '2' "
            "as the item_id — always use the actual 'id' value from the JSON above."
        ),
    }, indent=2)


@tool
def read_sharepoint_file(item_id: str) -> str:
    """
    Download a file from SharePoint by its Graph item ID and extract its text.

    Supports: PDF, DOCX, XLSX/CSV, TXT, MD.

    IMPORTANT: item_id must be the 'id' field from list_sharepoint_folder()
    results — a long alphanumeric string like '01ABCDEF123...'.
    Do NOT pass a display number like '1' or '2'.

    Returns the extracted text content plus file metadata — treat this exactly
    like the output of read_input_file() and use it to build the PRD.
    """
    try:
        import sharepoint_client as sp
    except ImportError as e:
        return json.dumps({"error": f"sharepoint_client import failed: {e}"})

    if not sp.is_configured():
        return json.dumps({
            "error": "SharePoint not configured.",
            "hint":  "Set SHAREPOINT_* env vars in SDLC/.env",
        })

    # ── Validate item_id — must be the real Graph ID, not a display number ───
    item_id = item_id.strip()
    if not item_id:
        return json.dumps({"error": "item_id is empty. Call list_sharepoint_folder() first."})

    if item_id.isdigit() or len(item_id) < 10:
        # Try to resolve from the cached folder listing
        if item_id.isdigit() and _sp_folder_cache:
            idx = int(item_id) - 1   # user picks are 1-based
            if 0 <= idx < len(_sp_folder_cache):
                item_id = _sp_folder_cache[idx]["id"]
                # fall through with the real ID
            else:
                return json.dumps({
                    "error": (
                        f"Display index '{item_id}' is out of range — "
                        f"the last listing had {len(_sp_folder_cache)} file(s). "
                        "Call list_sharepoint_folder() again to refresh."
                    ),
                })
        else:
            return json.dumps({
                "error": (
                    f"Invalid item_id '{item_id}'. "
                    "You must use the 'id' field from list_sharepoint_folder() — "
                    "a long alphanumeric string like '01J4NTCEB...'. "
                    "Call list_sharepoint_folder() first, then use the actual 'id' value."
                ),
            })

    # ── Download to a temp file ──────────────────────────────────────────────
    try:
        temp_path, filename = sp.download_to_temp(item_id)
    except Exception as e:
        return json.dumps({"error": f"Download failed: {e}"})

    # ── Extract text using the existing file readers ─────────────────────────
    suffix  = temp_path.suffix.lower()
    size_kb = round(temp_path.stat().st_size / 1024, 1)

    try:
        if suffix == ".pdf":
            content, ftype = _read_pdf(temp_path),  "PDF document"
        elif suffix == ".docx":
            content, ftype = _read_docx(temp_path), "Word document"
        elif suffix in (".xlsx", ".csv"):
            content, ftype = _read_csv(temp_path),  "Spreadsheet"
        elif suffix in (".txt", ".md", ".text"):
            content, ftype = _read_txt(temp_path),  "Text / Markdown"
        else:
            return json.dumps({
                "error":     f"Unsupported file type: {suffix}",
                "filename":  filename,
                "supported": [".pdf", ".docx", ".xlsx", ".csv", ".txt", ".md"],
            })
    finally:
        # Always clean up the temp file regardless of success/failure
        try:
            temp_path.unlink(missing_ok=True)
        except Exception:
            pass

    if not content.strip():
        return json.dumps({"error": "File appears to be empty or could not be parsed.", "filename": filename})

    truncated = False
    if len(content) > 32000:
        content   = content[:32000]
        truncated = True

    return json.dumps({
        "source":      "sharepoint",
        "filename":    filename,
        "file_type":   ftype,
        "size_kb":     size_kb,
        "truncated":   truncated,
        "content":     content,
        "instruction": (
            "Extract every PRD-relevant detail from this content: problem statement, "
            "business context, user types, requirements, constraints, stakeholders, "
            "goals, timelines, and open questions. "
            "Then identify ONLY what is missing and ask the user — one question at a time."
            + (" NOTE: content was truncated to 32 KB." if truncated else "")
        ),
    })


@tool
def upload_prd_to_sharepoint(md_path: str, folder_path: str) -> str:
    """
    Upload a locally-saved PRD Markdown file back to SharePoint.

    Args:
        md_path:     Absolute path to the .md file returned by save_prd()
                     (the md_path field in the save_prd result).
        folder_path: SharePoint site-relative folder to upload into,
                     e.g. "Documents/PRDs".
                     Defaults to SHAREPOINT_PRD_FOLDER env var if empty string.

    The file is uploaded with its original filename. If a file with the same
    name already exists in that SharePoint folder it is overwritten (Graph API
    PUT semantics).

    Returns the SharePoint web URL of the uploaded file.
    """
    try:
        import sharepoint_client as sp
    except ImportError as e:
        return json.dumps({"error": f"sharepoint_client import failed: {e}"})

    if not sp.is_configured():
        return json.dumps({
            "error": "SharePoint not configured.",
            "hint":  "Set SHAREPOINT_* env vars in SDLC/.env",
        })

    # ── Resolve folder ────────────────────────────────────────────────────────
    import os as _os
    target_folder = (
        folder_path.strip()
        or _os.getenv("SHAREPOINT_PRD_FOLDER", "")
        or _os.getenv("SHAREPOINT_READ_FOLDER", "")
        or ""
    )

    # ── Read local .md file ───────────────────────────────────────────────────
    local = Path(md_path.strip().strip('"').strip("'"))
    if not local.exists():
        return json.dumps({
            "error": f"Local file not found: {md_path}",
            "hint":  "Pass the md_path value returned by save_prd().",
        })

    try:
        content_bytes = local.read_bytes()
    except Exception as e:
        return json.dumps({"error": f"Could not read local file: {e}"})

    # ── Upload ────────────────────────────────────────────────────────────────
    try:
        result = sp.upload_file(
            folder_path  = target_folder,
            filename     = local.name,
            content      = content_bytes,
            content_type = "text/plain; charset=utf-8",
        )
    except Exception as e:
        return json.dumps({"error": f"SharePoint upload failed: {e}"})

    return json.dumps({
        "uploaded":    True,
        "filename":    result["name"],
        "web_url":     result["web_url"],
        "size_kb":     result["size_kb"],
        "sp_folder":   target_folder or "(root)",
        "local_file":  str(local),
    })


TOOLS = [
    get_current_date,
    load_product_context,
    read_input_file,
    list_existing_prds,
    load_prd,
    save_prd,
    list_sharepoint_folder,
    read_sharepoint_file,
    upload_prd_to_sharepoint,
]
